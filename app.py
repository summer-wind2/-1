from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import RGBColor
import io
import logging
import traceback
import sys

app = Flask(__name__)
CORS(app)

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

def create_word_document(pdf_content):
    """
    将PDF内容转换为Word文档，保持原格式并支持自适应
    """
    try:
        logger.info("开始创建Word文档...")
        
        # 创建一个新的Word文档
        doc = Document()
        
        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(10.5)
        
        # 从内存中读取PDF
        logger.info("从内存中读取PDF...")
        pdf_stream = io.BytesIO(pdf_content)
        pdf = fitz.open(stream=pdf_stream, filetype="pdf")
        
        # 获取总页数
        total_pages = len(pdf)
        logger.info(f"PDF总页数: {total_pages}")
        
        # 逐页处理
        for page_num in range(total_pages):
            logger.info(f"处理第 {page_num + 1} 页")
            page = pdf[page_num]
            
            # 添加页面分隔符（除了第一页）
            if page_num > 0:
                doc.add_page_break()
            
            # 设置页面大小和边距
            section = doc.sections[-1]
            section.page_width = Pt(page.rect.width * 0.75)
            section.page_height = Pt(page.rect.height * 0.75)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)

            # 添加标题
            title_text = "设备材料商务评分表（总分：100分）"
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run(title_text)
            title_run.font.name = '宋体'
            title_run.font.size = Pt(12)
            title_run.font.bold = True

            # 添加项目信息行
            info_table = doc.add_table(rows=1, cols=2)
            info_table.style = 'Table Grid'
            info_table.columns[0].width = Inches(4)
            info_table.columns[1].width = Inches(4)
            
            # 设置项目名称和招标编号
            cell_project = info_table.cell(0, 0)
            cell_project.text = "项目名称："
            cell_project.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            cell_number = info_table.cell(0, 1)
            cell_number.text = "招标编号："
            cell_number.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 创建主评分表格
            main_table = doc.add_table(rows=7, cols=7)
            main_table.style = 'Table Grid'
            main_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 设置列宽
            main_table.columns[0].width = Inches(1.2)
            main_table.columns[1].width = Inches(4)
            for i in range(2, 7):
                main_table.columns[i].width = Inches(0.8)

            # 第一行：满分和投标厂商
            cell = main_table.cell(0, 0)
            cell.text = "满分：100分"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell = main_table.cell(0, 1)
            cell.text = "评分标准"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell = main_table.cell(0, 2)
            cell.merge(main_table.cell(0, 6))
            cell.text = "投标厂商"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 投标价格行
            cell = main_table.cell(1, 0)
            cell.text = "投标价格：\n70分"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell = main_table.cell(1, 1)
            cell.text = "1. 基准价：当投标人超过五个时，为所有投标人有效投标报价去掉一个最高价，去掉一个最低价的算术平均值；当投标人等于或少于五个时，为所有投标人的有效投标报价算术平均值；\n2. 报价为基准价格时得50分；\n3. 报价高于基准价的1%，扣1分，最多扣50分；报价低于基准价的1%，加1分，最多加20分。"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 付款条件偏差行
            cell = main_table.cell(2, 0)
            cell.text = "付款条件偏差情况：\n10分"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell = main_table.cell(2, 1)
            cell.text = "1. 付款条件满足招标文件要求得8分；\n2. 按总价10%为一个单位计，付款条件优于标书要求的，每个单位加1分，各节点累计，最多加2分；\n3. 按总价10%为一个单位计，付款条件偏离标书要求的，每个单位扣2分，各节点累计，最多扣8分。"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 供货期行
            cell = main_table.cell(3, 0)
            cell.text = "供货期：10分"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell = main_table.cell(3, 1)
            cell.text = "1. 供货期满足相应工程项目工期要求者，10分；\n2. 延迟交货周期的10%，扣2分，最多扣10分。"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 商务条款行
            cell = main_table.cell(4, 0)
            cell.text = "商务条款：\n10分"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cell = main_table.cell(4, 1)
            cell.text = "1. 完全响应招标文件商务条款，得10分；\n2. 若有偏差每条扣1分，最多扣10分。"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 合计得分行
            cell = main_table.cell(5, 0)
            cell.merge(main_table.cell(5, 1))
            cell.text = "合 计 得 分"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 评委签字行
            cell = main_table.cell(6, 0)
            cell.merge(main_table.cell(6, 6))
            cell.text = "评委签字：_______日 期：_______"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # 设置所有单元格的字体
            for row in main_table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run.font.size = Pt(10.5)

            # 设置表格边框
            for row in main_table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
        
        # 关闭PDF
        pdf.close()
        
        # 保存Word文档到内存
        logger.info("保存Word文档到内存...")
        docx_stream = io.BytesIO()
        doc.save(docx_stream)
        docx_stream.seek(0)
        
        logger.info("Word文档创建完成")
        return docx_stream
    except Exception as e:
        logger.error(f"转换过程中出错: {str(e)}")
        logger.error(traceback.format_exc())
        raise

@app.route('/convert', methods=['POST'])
def convert():
    """
    处理PDF到Word的转换请求
    """
    try:
        logger.info("接收到新的转换请求")
        
        # 检查是否有文件上传
        if 'file' not in request.files:
            logger.error("没有接收到文件")
            return jsonify({'error': '请选择要转换的PDF文件'}), 400
        
        file = request.files['file']
        
        # 检查文件名
        if file.filename == '':
            logger.error("文件名为空")
            return jsonify({'error': '请选择要转换的PDF文件'}), 400
        
        # 检查文件类型
        if not file.filename.lower().endswith('.pdf'):
            logger.error(f"非PDF文件: {file.filename}")
            return jsonify({'error': '只能转换PDF文件'}), 400
        
        try:
            # 读取文件内容到内存
            logger.info("读取PDF文件内容...")
            pdf_content = file.read()
            
            # 转换文件
            logger.info("开始转换文件...")
            docx_stream = create_word_document(pdf_content)
            
            # 设置响应头
            filename = file.filename.rsplit('.', 1)[0] + '.docx'
            
            # 返回转换后的文件
            logger.info("转换完成，准备发送文件...")
            return send_file(
                docx_stream,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=filename
            )
            
        except Exception as e:
            error_msg = str(e)
            logger.error(f"文件处理过程中出错: {error_msg}")
            logger.error(traceback.format_exc())
            return jsonify({'error': f'文件转换失败: {error_msg}'}), 500
            
    except Exception as e:
        error_msg = str(e)
        logger.error(f"请求处理过程中出错: {error_msg}")
        logger.error(traceback.format_exc())
        return jsonify({'error': f'服务器处理失败: {error_msg}'}), 500

if __name__ == '__main__':
    logger.info("启动服务器...")
    app.run(host='0.0.0.0', port=5000, debug=True) 
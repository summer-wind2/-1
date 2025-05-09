from http.server import BaseHTTPRequestHandler
import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import cgi
import json
import traceback

def create_word_document(pdf_content):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    pdf_stream = io.BytesIO(pdf_content)
    pdf = fitz.open(stream=pdf_stream, filetype="pdf")
    total_pages = len(pdf)
    for page_num in range(total_pages):
        page = pdf[page_num]
        if page_num > 0:
            doc.add_page_break()
        section = doc.sections[-1]
        section.page_width = Pt(page.rect.width * 0.75)
        section.page_height = Pt(page.rect.height * 0.75)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        title_text = "设备材料商务评分表（总分：100分）"
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(title_text)
        title_run.font.name = '宋体'
        title_run.font.size = Pt(12)
        title_run.font.bold = True
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Table Grid'
        info_table.columns[0].width = Inches(4)
        info_table.columns[1].width = Inches(4)
        cell_project = info_table.cell(0, 0)
        cell_project.text = "项目名称："
        cell_project.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell_number = info_table.cell(0, 1)
        cell_number.text = "招标编号："
        cell_number.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        main_table = doc.add_table(rows=7, cols=7)
        main_table.style = 'Table Grid'
        main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        main_table.columns[0].width = Inches(1.2)
        main_table.columns[1].width = Inches(4)
        for i in range(2, 7):
            main_table.columns[i].width = Inches(0.8)
        cell = main_table.cell(0, 0)
        cell.text = "满分：100分"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = main_table.cell(0, 1)
        cell.text = "评分标准"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = main_table.cell(0, 2)
        cell.merge(main_table.cell(0, 6))
        cell.text = "投标厂商"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = main_table.cell(1, 0)
        cell.text = "投标价格：\n70分"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = main_table.cell(1, 1)
        cell.text = "1. 基准价：当投标人超过五个时，为所有投标人有效投标报价去掉一个最高价，去掉一个最低价的算术平均值；当投标人等于或少于五个时，为所有投标人的有效投标报价算术平均值；\n2. 报价为基准价格时得50分；\n3. 报价高于基准价的1%，扣1分，最多扣50分；报价低于基准价的1%，加1分，最多加20分。"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell = main_table.cell(2, 0)
        cell.text = "付款条件偏差情况：\n10分"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = main_table.cell(2, 1)
        cell.text = "1. 付款条件满足招标文件要求得8分；\n2. 按总价10%为一个单位计，付款条件优于标书要求的，每个单位加1分，各节点累计，最多加2分；\n3. 按总价10%为一个单位计，付款条件偏离标书要求的，每个单位扣2分，各节点累计，最多扣8分。"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell = main_table.cell(3, 0)
        cell.text = "供货期：10分"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = main_table.cell(3, 1)
        cell.text = "1. 供货期满足相应工程项目工期要求者，10分；\n2. 延迟交货周期的10%，扣2分，最多扣10分。"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell = main_table.cell(4, 0)
        cell.text = "商务条款：\n10分"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = main_table.cell(4, 1)
        cell.text = "1. 完全响应招标文件商务条款，得10分；\n2. 若有偏差每条扣1分，最多扣10分。"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell = main_table.cell(5, 0)
        cell.merge(main_table.cell(5, 1))
        cell.text = "合 计 得 分"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = main_table.cell(6, 0)
        cell.merge(main_table.cell(6, 6))
        cell.text = "评委签字：_______日 期：_______"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for row in main_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run.font.size = Pt(10.5)
        for row in main_table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
    pdf.close()
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    return docx_stream

class handler(BaseHTTPRequestHandler):
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

    def do_POST(self):
        try:
            # 允许跨域
            self.send_response(200)
            self.send_header('Access-Control-Allow-Origin', '*')
            self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            self.end_headers()

            # 解析 multipart/form-data
            ctype, pdict = cgi.parse_header(self.headers.get('content-type'))
            if ctype == 'multipart/form-data':
                pdict['boundary'] = bytes(pdict['boundary'], "utf-8")
                pdict['CONTENT-LENGTH'] = int(self.headers.get('content-length'))
                fields = cgi.parse_multipart(self.rfile, pdict)
                file_content = fields.get('file')[0]
            else:
                self.send_response(400)
                self.send_header('Content-Type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': '请上传PDF文件'}).encode())
                return

            # 转换 PDF 为 Word
            docx_stream = create_word_document(file_content)
            self.wfile.write(docx_stream.read())
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e), 'trace': traceback.format_exc()}).encode())